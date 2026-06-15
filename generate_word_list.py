# -*- coding: utf-8 -*-
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 상수 ───────────────────────────────────────────────────────────────────
FONT_KO = "맑은 고딕"

BUILTIN_GESTURES = [
    ("Thumb_Up",    "엄지 위로",     "좋아요",     "엄지손가락을 위로 세우는 동작"),
    ("Thumb_Down",  "엄지 아래로",   "아니오",     "엄지손가락을 아래로 향하는 동작"),
    ("Open_Palm",   "손바닥 펼치기", "안녕하세요", "다섯 손가락을 모두 펼친 상태"),
    ("Closed_Fist", "주먹 쥐기",     "잠깐만요",   "다섯 손가락을 모두 오므린 상태"),
    ("Victory",     "V자 (승리)",    "감사합니다", "검지·중지를 펴 V 모양"),
    ("Pointing_Up", "위로 가리키기", "주세요",     "검지손가락만 위로 세우는 동작"),
    ("ILoveYou",    "I Love You",    "도와주세요", "엄지·검지·새끼를 편 ASL 수형"),
]

WORD_CATEGORIES = {
    "공통기본": [
        "나", "너", "우리", "이것", "저것", "여기", "거기", "있다", "없다",
        "하다", "가다", "오다", "보다", "먹다", "마시다", "자다", "일어나다",
        "좋다", "나쁘다", "크다", "작다", "많다", "적다", "빠르다", "느리다",
        "쉽다", "어렵다", "모르다", "이해", "얼마",
    ],
    "일상동사": [
        "걷다", "달리다", "서다", "앉다", "듣다", "말하다", "읽다", "쓰다",
        "놀다", "일하다", "요리하다", "청소하다", "씻다", "입다", "사다",
        "팔다", "열다", "닫다", "출근", "훈련",
    ],
    "감정·상태": [
        "기쁘다", "슬프다", "화나다", "놀라다", "무섭다", "피곤하다",
        "힘들다", "아프다", "졸리다", "배고프다", "배부르다", "목마르다",
        "따뜻하다", "덥다", "춥다", "시원하다", "바쁘다", "조용하다",
    ],
    "장소·교통": [
        "병원", "학교", "집", "회사", "식당", "마트", "공원",
        "기차", "기차역", "구급차", "주차장",
        "왼쪽", "오른쪽", "위", "아래", "지도",
    ],
    "사람·관계": [
        "나", "너", "가족", "남자", "여자", "노인", "남편", "아내",
        "딸", "친구", "의사", "간호사", "나이", "여동생", "할머니", "할아버지",
    ],
    "식사·쇼핑": [
        "밥", "우유", "커피", "달다", "식당",
    ],
    "긴급·중요": [
        "급하다", "구조", "조심", "경찰서",
        "통증", "골절", "입원", "퇴원", "검사", "상담",
    ],
    "직장·학교": [
        "시험", "쓰다", "성공", "계획", "평가", "동아리", "출근", "훈련",
    ],
    "색깔": [
        "빨강", "노랑", "초록", "파랑", "검정",
    ],
}

TRAINED_WORDS = {
    "가다", "가족", "간호사", "걷다", "검사", "검정", "경찰서", "계획",
    "골절", "구급차", "구조", "급하다", "기차", "기차역", "나", "나쁘다",
    "나이", "남자", "남편", "너", "노랑", "노인", "놀다", "놀라다",
    "느리다", "달다", "달리다", "덥다", "동아리", "듣다", "따뜻하다",
    "딸", "모르다", "목마르다", "무섭다", "바쁘다", "밥", "배고프다",
    "배부르다", "병원", "빠르다", "빨강", "빨리", "상담", "서다", "성공",
    "쉽다", "슬프다", "시원하다", "시험", "식당", "쓰다", "아내", "아프다",
    "어렵다", "얼마", "여동생", "여자", "오다", "오른쪽", "왼쪽", "우유",
    "위", "의사", "이해", "일어나다", "입원", "자다", "조심", "조용하다",
    "졸리다", "좋다", "주차장", "지도", "초록", "출근", "춥다", "친구",
    "커피", "통증", "퇴원", "파랑", "평가", "피곤하다", "할머니", "할아버지",
    "화나다", "훈련", "힘들다",
}

C_HEADER = "1F497D"   # 진한 파랑 (헤더)
C_SEC    = "2E74B5"   # 중간 파랑 (섹션 제목)
C_TRN    = "005C00"   # 진한 초록 (학습 완료)
C_UNTR   = "CC0000"   # 빨강 (미학습)
C_BG_TRN = "E2EFDA"   # 연한 초록 배경
C_BG_UNT = "FFE7E7"   # 연한 빨강 배경
C_WHITE  = "FFFFFF"
C_GRAY   = "909090"


# ── 헬퍼 ───────────────────────────────────────────────────────────────────
def _set_font(run, bold=False, size=10, hex_color=None):
    """맑은 고딕 + 색상 + 크기 한번에 설정"""
    run.font.name = FONT_KO
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    run.bold = bold
    run.font.size = Pt(size)
    if hex_color:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _heading(doc, text, size=14, hex_color=C_SEC, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    _set_font(r, bold=True, size=size, hex_color=hex_color)
    return p


def _note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=9, hex_color="404040")
    p.paragraph_format.space_after = Pt(6)
    return p


# ── 문서 본문 ──────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # 기본 폰트를 맑은 고딕으로 변경
    doc.styles["Normal"].font.name = FONT_KO
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)

    # 페이지 여백
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── 제목 ──────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run("수화 인식 서비스 — 제스처 & 단어 목록")
    _set_font(tr, bold=True, size=18, hex_color=C_HEADER)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(12)
    sr = sp.add_run("Korean Sign Language Recognition Service  |  2026.05")
    _set_font(sr, size=10, hex_color=C_GRAY)

    # ── 범례 ──────────────────────────────────────────────
    lp = doc.add_paragraph()
    lp.paragraph_format.space_after = Pt(10)
    lr = lp.add_run("범례:  ")
    _set_font(lr, bold=True, size=10)

    g1 = lp.add_run("■ 학습 완료 (모델에 포함)")
    _set_font(g1, bold=True, size=10, hex_color=C_TRN)

    lp.add_run("     ")

    g2 = lp.add_run("■ 미학습 (데이터 수집 필요)")
    _set_font(g2, bold=True, size=10, hex_color=C_UNTR)

    # ══════════════════════════════════════════════════════
    # 섹션 1: MediaPipe 내장 제스처
    # ══════════════════════════════════════════════════════
    _heading(doc, "1. MediaPipe 내장 제스처 (7개)", size=14, space_before=10)
    _note(doc, "아래 7가지 손 모양은 별도의 데이터 수집 없이 MediaPipe GestureRecognizer 모델로 즉시 인식됩니다.")

    tbl1 = doc.add_table(rows=1, cols=4)
    tbl1.style = "Table Grid"
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w1 = [Cm(4.0), Cm(3.5), Cm(3.0), Cm(8.0)]
    hdrs1  = ["제스처 이름", "손 모양 설명", "한국어 단어", "비고"]
    hdr_cells = tbl1.rows[0].cells
    for i, (cell, h) in enumerate(zip(hdr_cells, hdrs1)):
        _set_cell_bg(cell, C_HEADER)
        cell.width = col_w1[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        _set_font(run, bold=True, size=10, hex_color=C_WHITE)

    for gesture, shape, word, desc in BUILTIN_GESTURES:
        row = tbl1.add_row().cells
        vals = [gesture, shape, word, desc]
        for i, (cell, val) in enumerate(zip(row, vals)):
            cell.width = col_w1[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(val)
            if i == 2:
                _set_font(run, bold=True, size=10, hex_color=C_TRN)
            else:
                _set_font(run, size=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 섹션 2: AIHub KSL 학습 모델 단어
    # ══════════════════════════════════════════════════════
    _heading(doc, f"2. AIHub KSL 학습 모델 단어 (총 {len(TRAINED_WORDS)}개)", size=14, space_before=10)
    _note(doc,
        "한국 수어(KSL) 표준 수형 기반으로 학습된 단어입니다. "
        "초록 배경 = 모델 학습 완료, 빨간 배경 = 미학습.\n"
        "* 현재 모델은 AIHub 표준 수형 기반이므로 개인 인식률을 높이려면 직접 데이터 수집이 필요합니다.")

    for cat_name, words in WORD_CATEGORIES.items():
        _heading(doc, f"  ▸ {cat_name}", size=11, space_before=8)

        COLS = 5
        rows_n = (len(words) + COLS - 1) // COLS
        tbl2 = doc.add_table(rows=rows_n, cols=COLS)
        tbl2.style = "Table Grid"
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, word in enumerate(sorted(words)):
            r = idx // COLS
            c = idx % COLS
            cell = tbl2.cell(r, c)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(word)
            if word in TRAINED_WORDS:
                _set_font(run, bold=True, size=10, hex_color=C_TRN)
                _set_cell_bg(cell, C_BG_TRN)
            else:
                _set_font(run, size=10, hex_color=C_UNTR)
                _set_cell_bg(cell, C_BG_UNT)

        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 섹션 3: 요약
    # ══════════════════════════════════════════════════════
    _heading(doc, "3. 요약 통계", size=14, space_before=10)

    total_cat = sum(len(v) for v in WORD_CATEGORIES.values())
    matched   = sum(1 for v in WORD_CATEGORIES.values() for w in v if w in TRAINED_WORDS)

    stats = [
        ("MediaPipe 즉시 인식 가능 제스처",       "7개",                    "별도 학습 불필요"),
        ("AIHub KSL 모델 학습 완료 단어",         f"{len(TRAINED_WORDS)}개", "현재 모델에 포함"),
        ("카테고리 내 학습 완료 (중복 포함)",      f"{matched}개",            f"/ {total_cat}개 전체"),
        ("추가 데이터 수집 시 인식 목표 단어 수",  f"~{total_cat}개+",        "개인 데이터 수집 후"),
    ]

    tbl3 = doc.add_table(rows=len(stats), cols=3)
    tbl3.style = "Table Grid"
    col_w3 = [Cm(9.5), Cm(2.5), Cm(6.5)]
    for i, (label, val, note) in enumerate(stats):
        cells = tbl3.rows[i].cells
        for j, (cell, txt) in enumerate(zip(cells, [label, val, note])):
            cell.width = col_w3[j]
            para = cell.paragraphs[0]
            run  = para.add_run(txt)
            if j == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_font(run, bold=True, size=10, hex_color=C_TRN)
            else:
                _set_font(run, size=10)

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("생성일: 2026-05-29  |  Korean Sign Language Recognition Service")
    _set_font(fr, size=8, hex_color=C_GRAY)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "수화_단어_목록.docx")
    doc.save(out)
    print("saved:", out)
    return out


if __name__ == "__main__":
    main()
